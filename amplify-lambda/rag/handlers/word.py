# Copyright (c) 2024 Vanderbilt University
# Authors: Jules White, Allen Karns, Karely Rodriguez, Max Moundas

from docx import Document
import io
from enum import Enum
from rag.handlers.text import TextExtractionHandler
from rag.handlers.shared_functions import (
    hash_visual_data,
    ensure_supported_format,
    format_visual_chunk_data,
)


PNG = "image/png"


class VisualType(Enum):
    """Enum for visual content types we extract from Word documents"""

    IMAGE = "Image"  # Photos, screenshots, logos, imported images


class DOCXHandler(TextExtractionHandler):
    def extract_text(self, file_content, visual_map={}):
        """
        Extract text and visual content from Word document.
        Now supports visual_map from preprocessing for multimodal content.
        """
        with io.BytesIO(file_content) as f:
            doc = Document(f)

            # Preprocess visual_map to group visuals by approximate location for efficient lookup
            visuals_by_section = {}
            visuals_by_paragraph = {}
            floating_visuals = []  # Visuals without clear paragraph association

            for visual_marker, visual_data in visual_map.items():
                location = visual_data.get("location", {})
                section_number = location.get("section_number")
                paragraph_number = location.get("paragraph_number")

                if paragraph_number:
                    # Visual tied to specific paragraph
                    if paragraph_number not in visuals_by_paragraph:
                        visuals_by_paragraph[paragraph_number] = []
                    visuals_by_paragraph[paragraph_number].append(visual_data)
                elif section_number:
                    # Visual tied to section but not specific paragraph
                    if section_number not in visuals_by_section:
                        visuals_by_section[section_number] = []
                    visuals_by_section[section_number].append(visual_data)
                else:
                    # Floating visual - we'll add at end
                    floating_visuals.append(visual_data)

            # Use centralized section mapping
            section_mapping = self._build_section_mapping(doc)
            chunks = []

            # Create an array to hold all the text and structure information
            for i, paragraph in enumerate(doc.paragraphs, start=1):
                # Only add non-empty paragraphs
                if paragraph.text.strip():
                    section_info = section_mapping.get(i, {})
                    paragraph_info = {
                        "content": paragraph.text,
                        "tokens": self.num_tokens_from_string(paragraph.text),
                        "location": {
                            "section_number": section_info.get("section_number", 1),
                            "paragraph_number": i,
                            "section_title": section_info.get("section_title", ""),
                        },
                        "canSplit": True,
                    }
                    chunks.append(paragraph_info)

                # Check for visuals associated with this specific paragraph
                paragraph_visuals = visuals_by_paragraph.get(i, [])
                for visual_data in paragraph_visuals:
                    if visual_data.get("transcription"):
                        visual_chunk = format_visual_chunk_data(visual_data, self.num_tokens_from_string)
                        chunks.append(visual_chunk)

                # If this is the last paragraph of a section, add section-level visuals
                next_paragraph_is_new_section = False
                if i < len(doc.paragraphs):
                    next_paragraph = doc.paragraphs[
                        i
                    ]  # i is 1-indexed, so this gets next paragraph
                    if next_paragraph.style.name.startswith("Heading"):
                        next_paragraph_is_new_section = True

                # Add section visuals at end of section or end of document
                if next_paragraph_is_new_section or i == len(doc.paragraphs):
                    current_section = section_mapping.get(i, {}).get(
                        "section_number", 1
                    )
                    section_visuals = visuals_by_section.get(current_section, [])
                    for visual_data in section_visuals:
                        if visual_data.get("transcription"):
                            visual_chunk = format_visual_chunk_data(visual_data, self.num_tokens_from_string)
                            chunks.append(visual_chunk)

            # Add floating visuals at the end
            for visual_data in floating_visuals:
                if visual_data.get("transcription"):
                    visual_chunk = format_visual_chunk_data(visual_data, self.num_tokens_from_string)
                    chunks.append(visual_chunk)

            return chunks

    ### Visual Data Extraction ###
    def preprocess_docx_visuals(self, file_content):
        """
        Extract visual content, inject visual markers, and remove images from document.
        Returns modified document content with visual markers for MarkItDown processing.
        """

        with io.BytesIO(file_content) as f:
            doc = Document(f)

        visual_map = {}

        # Build section mapping first for location context
        section_mapping = self._build_section_mapping(doc)

        # Extract visuals from the document
        document_visuals = self.extract_document_visuals(doc, section_mapping)

        # Add visuals to global map
        visual_map.update(document_visuals)

        # Inject visual markers and remove images from the document
        self._inject_visual_markers_and_remove_images(doc, visual_map)

        # Save the modified document to bytes
        modified_content = self._document_to_bytes(doc)

        return modified_content, visual_map

    def extract_document_visuals(self, doc, section_mapping):
        """Extract all visual content from Word document"""

        visuals = {}
        # Count each visual type we extract
        type_counters = {
            VisualType.IMAGE.value: 0,
        }

        # Extract from comments and annotations
        comment_visuals = self.extract_comment_visuals(doc, type_counters)
        visuals.update(comment_visuals)

        # Extract from inline images in paragraphs
        for paragraph_idx, paragraph in enumerate(doc.paragraphs, start=1):
            for run in paragraph.runs:
                # Check for images
                pic_elements = run._element.xpath(".//pic:pic")
                if pic_elements:
                    for pic_element in pic_elements:
                        type_counters[VisualType.IMAGE.value] += 1
                        marker = f"<Visual#{paragraph_idx}_{VisualType.IMAGE.value}_{type_counters[VisualType.IMAGE.value]}>"
                        
                        section_info = section_mapping.get(paragraph_idx, {})
                        visual_data = self.extract_inline_image_data(
                            pic_element, paragraph_idx, section_info, doc
                        )
                        if visual_data:
                            visuals[marker] = visual_data

        # Extract from headers/footers if present
        for section_idx, section in enumerate(doc.sections):
            try:
                if section.header:
                    header_visuals = self.extract_header_footer_visuals(
                        section.header, "header", type_counters, doc
                    )
                    visuals.update(header_visuals)

                if section.footer:
                    footer_visuals = self.extract_header_footer_visuals(
                        section.footer, "footer", type_counters, doc
                    )
                    visuals.update(footer_visuals)

            except Exception as e:
                print(
                    f"[DEBUG] Error processing header/footer for section {section_idx + 1}: {e}"
                )
                continue

        return visuals

    def extract_inline_image_data(
        self, pic_element, paragraph_number, section_info, doc
    ):
        """Extract data from inline image in paragraph"""
        # Use centralized image extraction pipeline
        image_data = self._extract_image_from_relationship(
            pic_element, doc, f"paragraph {paragraph_number}"
        )

        if not image_data:
            return None

        # Create location info
        location_info = {
            "section_number": section_info.get("section_number", 1),
            "paragraph_number": paragraph_number,
            "section_title": section_info.get("section_title", ""),
            "position_type": "inline",
        }

        # Use centralized visual data structure creation
        return self._create_visual_data_structure(
            VisualType.IMAGE.value, image_data, location_info
        )

    def _find_document_part_from_element(self, element):
        """Try to find the document part from an XML element"""
        try:
            # Look for common parent elements that might contain relationships
            current = element
            while current is not None:
                # Check if this element has relationships
                if hasattr(current, "rels"):
                    return current

                # Check if this element has a part attribute
                if hasattr(current, "part"):
                    return current.part

                # Move up the parent chain
                current = current.getparent()

                # Safety check to prevent infinite loops
                if current is element:
                    break

        except Exception as e:
            print(f"Error in _find_document_part_from_element: {e}")

        return None

    def _find_rels_through_parent_chain(self, element):
        """Try to find relationships by traversing the parent chain"""
        try:
            current = element
            max_depth = 10  # Prevent infinite loops

            for _ in range(max_depth):
                if current is None:
                    break

                # Check if this element has relationships
                if hasattr(current, "rels"):
                    return current

                # Check if this element has a part attribute
                if hasattr(current, "part"):
                    part = current.part
                    if hasattr(part, "rels"):
                        return part

                # Move up the parent chain
                current = current.getparent()

        except Exception as e:
            print(f"Error in _find_rels_through_parent_chain: {e}")

        return None

    def extract_header_footer_visuals(
        self, header_footer, location_type, type_counters, doc
    ):
        """Extract visuals from headers/footers"""
        visuals = {}
        processed_hashes = (
            set()
        )  # Track processed images by their hash to avoid duplicates

        try:
            # Method 1: Look for images in header/footer paragraphs (existing approach)
            for paragraph_idx, paragraph in enumerate(header_footer.paragraphs):
                for run_idx, run in enumerate(paragraph.runs):
                    pic_elements = run._element.xpath(".//pic:pic")
                    if pic_elements:
                        for pic_idx, pic_element in enumerate(pic_elements):
                            # Extract the image first to get its hash
                            image_data = self._extract_image_from_relationship(
                                pic_element,
                                doc,
                                f"{location_type} paragraph {paragraph_idx + 1}",
                            )

                            if (
                                image_data
                                and image_data["hash"] not in processed_hashes
                            ):
                                processed_hashes.add(
                                    image_data["hash"]
                                )  # Mark this image as processed
                                type_counters[VisualType.IMAGE.value] += 1
                                marker = f"<Visual#{location_type}_{VisualType.IMAGE.value}_{type_counters[VisualType.IMAGE.value]}>"

                                # Create location info
                                location_info = {
                                    "section_number": 0,  # Header/footer spans document
                                    "paragraph_number": None,
                                    "position_type": location_type,
                                }

                                # Use centralized visual data structure creation
                                visual_data = self._create_visual_data_structure(
                                    VisualType.IMAGE.value, image_data, location_info
                                )

                                # Override title for header/footer images
                                if (
                                    not visual_data["title"]
                                    or visual_data["title"] == VisualType.IMAGE.value
                                ):
                                    visual_data["title"] = (
                                        f"{location_type.title()} Image"
                                    )

                                visuals[marker] = visual_data

            # Method 2: Look for images directly in the header/footer XML structure
            print(f"[DEBUG] Checking {location_type} XML structure for images...")
            try:
                # Get the XML element of the header/footer
                header_footer_element = header_footer._element

                # Look for images in the entire header/footer structure
                all_pic_elements = header_footer_element.xpath(".//pic:pic")

                for pic_idx, pic_element in enumerate(all_pic_elements):
                    # Extract the image first to check its hash
                    image_data = self._extract_image_from_relationship(
                        pic_element, doc, f"{location_type} XML structure"
                    )

                    # Only process if it's a new image
                    if image_data and image_data["hash"] not in processed_hashes:
                        processed_hashes.add(
                            image_data["hash"]
                        )  # Mark this image as processed
                        type_counters[VisualType.IMAGE.value] += 1
                        marker = f"<Visual#{location_type}_{VisualType.IMAGE.value}_{type_counters[VisualType.IMAGE.value]}>"

                        # Create location info
                        location_info = {
                            "section_number": 0,  # Header/footer spans document
                            "paragraph_number": None,
                            "section_title": f"Document {location_type.title()}",
                            "position_type": location_type,
                        }

                        # Use centralized visual data structure creation
                        visual_data = self._create_visual_data_structure(
                            VisualType.IMAGE.value, image_data, location_info
                        )

                        # Override title for header/footer images
                        if (
                            not visual_data["title"]
                            or visual_data["title"] == VisualType.IMAGE.value
                        ):
                            visual_data["title"] = f"{location_type.title()} Image"

                        visuals[marker] = visual_data

            except Exception as e:
                print(f"[DEBUG] Error in Method 2 for {location_type}: {e}")

            print(f"[DEBUG] {location_type} extraction complete. Found {len(visuals)} unique visuals")

        except Exception as e:
            print(f"Error extracting {location_type} visuals: {e}")

        return visuals

    def extract_pic_metadata(self, pic_element):
        """Extract metadata from picture element"""
        metadata = {}

        try:
            # Try to get alt text and title from XML
            metadata["alt_text"] = ""
            metadata["title"] = ""
            metadata["hyperlink"] = ""

            # Look for description and title in the XML structure
            # This is simplified - real implementation would parse XML more thoroughly
            desc_elements = pic_element.xpath(".//pic:cNvPr/@descr")
            if desc_elements:
                metadata["alt_text"] = desc_elements[0]

            name_elements = pic_element.xpath(".//pic:cNvPr/@name")
            if name_elements:
                metadata["title"] = name_elements[0]

        except Exception as e:
            print(f"Error extracting pic metadata: {e}")

        return metadata

    def _find_document_part_with_relationships(self, element, doc, context_name=""):
        """Centralized method to find document part with relationships using 4-tier approach"""
        doc_part = None

        # Approach 1: Try to navigate up the XML tree
        try:
            root = element.getroottree()
            if root is not None:
                root_element = root.getroot()
                if root_element is not None:
                    doc_part = root_element.getparent()
                    # Navigate to find the part with relationships
                    while doc_part is not None and not hasattr(doc_part, "rels"):
                        doc_part = doc_part.getparent()
        except Exception as e:
            print(f"{context_name} approach 1 failed: {e}")
            doc_part = None

        # Approach 2: Try to find it through the element's namespace
        if doc_part is None or not hasattr(doc_part, "rels"):
            try:
                doc_part = self._find_document_part_from_element(element)
            except Exception as e:
                print(f"{context_name} approach 2 failed: {e}")
                doc_part = None

        # Approach 3: Try to find relationships through the element's direct parent chain
        if doc_part is None or not hasattr(doc_part, "rels"):
            try:
                doc_part = self._find_rels_through_parent_chain(element)
            except Exception as e:
                print(f"{context_name} approach 3 failed: {e}")
                doc_part = None

        # Approach 4: Use the document object directly (most reliable)
        if (
            doc_part is None
            or not hasattr(doc, "part")
            and not hasattr(doc.part, "rels")
        ):
            try:
                if hasattr(doc, "part") and hasattr(doc.part, "rels"):
                    doc_part = doc.part
            except Exception as e:
                print(f"{context_name} approach 4 failed: {e}")
                doc_part = None

        return doc_part

    def _extract_image_from_relationship(self, pic_element, doc, context_name=""):
        """Centralized image extraction pipeline"""
        try:
            # Get the relationship ID for the image
            blip_elements = pic_element.xpath(".//a:blip/@r:embed")
            if not blip_elements:
                print(f"No image relationship found in {context_name}")
                return None

            rId = blip_elements[0]

            # Find document part with relationships
            doc_part = self._find_document_part_with_relationships(
                pic_element, doc, context_name
            )

            if doc_part and hasattr(doc_part, "rels") and rId in doc_part.rels:
                rel = doc_part.rels[rId]
                image_part = rel.target_part
                image_bytes = image_part.blob
                original_format = image_part.content_type

                # Convert format if necessary
                final_image_bytes, final_format = ensure_supported_format(
                    image_bytes, original_format
                )

                # Generate hash for deduplication
                content_hash = hash_visual_data(final_image_bytes)

                # Extract metadata
                metadata = self.extract_pic_metadata(pic_element)

                return {
                    "image_bytes": final_image_bytes,
                    "format": final_format,
                    "hash": content_hash,
                    "metadata": metadata,
                    "original_format": original_format,
                }
            else:
                print(f"Could not find image relationship {rId} in {context_name}")
                return None

        except Exception as e:
            print(f"Image extraction failed in {context_name}: {e}")
            return None

    def _create_visual_data_structure(
        self, visual_type, image_data, location_info, text_content=""
    ):
        """Centralized visual data structure creation"""
        return {
            "type": visual_type,
            "format": image_data["format"],
            "data": image_data["image_bytes"],
            "hash": image_data["hash"],
            "location": location_info,
            "alt_text": image_data["metadata"].get("alt_text", ""),
            "title": image_data["metadata"].get("title", ""),
            "hyperlink": image_data["metadata"].get("hyperlink", ""),
            "text_content": text_content,
            "original_format": image_data.get("original_format", ""),
        }

    def _build_section_mapping(self, doc):
        """Centralized section mapping logic"""
        section_mapping = {}
        current_section_index = 0
        section_headers = {}

        for i, paragraph in enumerate(doc.paragraphs, start=1):
            if paragraph.style.name.startswith("Heading"):
                current_section_index += 1
                section_headers[current_section_index] = paragraph.text

            section_mapping[i] = {
                "section_number": current_section_index,
                "section_title": section_headers.get(current_section_index, ""),
            }

        return section_mapping

    def extract_comment_visuals(self, doc, type_counters):
        """Extract visuals from comments and annotations"""
        visuals = {}

        try:
            # Get comments from the document
            if not hasattr(doc, "part") or not hasattr(doc.part, "rels"):
                return visuals

            # Look for comment relationships
            for rel_id, rel in doc.part.rels.items():
                if "comments" in rel.reltype.lower():
                    try:
                        comments_part = rel.target_part
                        if not hasattr(comments_part, "_element"):
                            continue

                        # Look for images in comments
                        comment_elements = comments_part._element.xpath(".//w:comment")

                        for comment_idx, comment in enumerate(comment_elements):
                            try:
                                # Look for images in this comment
                                pic_elements = comment.xpath(".//pic:pic")
                                if pic_elements:
                                    for pic_idx, pic_element in enumerate(pic_elements):
                                        type_counters[VisualType.IMAGE.value] += 1
                                        marker = f"<Visual#Comment_{VisualType.IMAGE.value}_{type_counters[VisualType.IMAGE.value]}>"

                                        # Extract the image
                                        image_data = (
                                            self._extract_image_from_relationship(
                                                pic_element,
                                                doc,
                                                f"comment {comment_idx + 1}",
                                            )
                                        )

                                        if image_data:
                                            # Get comment text for context
                                            comment_text = " ".join(
                                                text for text in comment.itertext()
                                            ).strip()

                                            # Create location info
                                            location_info = {
                                                "section_number": 0,  # Comments span document
                                                "paragraph_number": None,
                                                "section_title": "Document Comments",
                                                "position_type": "comment",
                                                "comment_text": (
                                                    comment_text[:100] + "..."
                                                    if len(comment_text) > 100
                                                    else comment_text
                                                ),
                                            }

                                            # Create visual data structure
                                            visual_data = (
                                                self._create_visual_data_structure(
                                                    VisualType.IMAGE.value,
                                                    image_data,
                                                    location_info,
                                                    comment_text,
                                                )
                                            )

                                            # Override title
                                            if (
                                                not visual_data["title"]
                                                or visual_data["title"]
                                                == VisualType.IMAGE.value
                                            ):
                                                visual_data["title"] = (
                                                    f"Comment Image {comment_idx + 1}"
                                                )

                                            visuals[marker] = visual_data

                            except Exception as e:
                                print( f"[DEBUG] Error processing comment {comment_idx + 1}: {e}" )
                                continue

                    except Exception as e:
                        print(f"[DEBUG] Error processing comments part: {e}")
                        continue

        except Exception as e:
            print(f"[DEBUG] Error extracting comment visuals: {e}")

        return visuals

    def _inject_visual_markers_and_remove_images(self, doc, visual_map):
        """
        Inject visual markers and remove images from the document.
        Processes paragraphs, headers, and footers.
        """
        
        # Create a reverse mapping from location to marker for efficient lookup
        location_to_marker = {}
        for marker, visual_data in visual_map.items():
            location = visual_data.get("location", {})
            paragraph_number = location.get("paragraph_number")
            position_type = location.get("position_type", "inline")
            
            if paragraph_number:
                # Regular paragraph image
                key = f"paragraph_{paragraph_number}"
                if key not in location_to_marker:
                    location_to_marker[key] = []
                location_to_marker[key].append(marker)
            elif position_type in ["header", "footer"]:
                # Header/footer image
                key = f"{position_type}_image"
                if key not in location_to_marker:
                    location_to_marker[key] = []
                location_to_marker[key].append(marker)
        
        # Process main document paragraphs
        for paragraph_idx, paragraph in enumerate(doc.paragraphs, start=1):
            key = f"paragraph_{paragraph_idx}"
            if key in location_to_marker:
                markers = location_to_marker[key]
                self._replace_images_with_markers_in_paragraph(paragraph, markers)
        
        # Process headers and footers
        for section in doc.sections:
            if section.header:
                header_markers = location_to_marker.get("header_image", [])
                if header_markers:
                    for paragraph in section.header.paragraphs:
                        self._replace_images_with_markers_in_paragraph(paragraph, header_markers)
            
            if section.footer:
                footer_markers = location_to_marker.get("footer_image", [])
                if footer_markers:
                    for paragraph in section.footer.paragraphs:
                        self._replace_images_with_markers_in_paragraph(paragraph, footer_markers)
        
        print(f"[DEBUG] Finished injecting visual markers for {len(visual_map)} visuals")

    def _replace_images_with_markers_in_paragraph(self, paragraph, markers):
        """
        Replace image elements with visual markers in a specific paragraph.
        """
        marker_index = 0
        
        for run in paragraph.runs:
            # Find all image elements in this run
            pic_elements = run._element.xpath(".//pic:pic")
            
            if pic_elements:
                # Remove all image elements from this run
                for pic_element in pic_elements:
                    # Remove the entire drawing element that contains the image
                    drawing_element = pic_element
                    while drawing_element is not None and drawing_element.tag.split('}')[-1] != 'drawing':
                        drawing_element = drawing_element.getparent()
                    
                    if drawing_element is not None:
                        drawing_element.getparent().remove(drawing_element)
                
                # Add visual marker text to the run
                if marker_index < len(markers):
                    # Use the full marker directly
                    run.text = markers[marker_index]
                    marker_index += 1
                else:
                    # Fallback if we run out of markers
                    run.text = "<Visual#Unknown>"
        
        return marker_index

    def _document_to_bytes(self, doc):
        """Converts a Document object to bytes."""
        with io.BytesIO() as output:
            doc.save(output)
            return output.getvalue()
