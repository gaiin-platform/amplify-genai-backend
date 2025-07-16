# Copyright (c) 2024 Vanderbilt University
# Authors: Jules White, Allen Karns, Karely Rodriguez, Max Moundas

from docx import Document
import io

from rag.handlers.text import TextExtractionHandler


class DOCXHandler(TextExtractionHandler):
    def extract_text(self, file_content, encoding):
        with io.BytesIO(file_content) as f:
            doc = Document(f)
            # Create a structure to hold section header information
            current_section_index = 0
            section_headers = {}

            # Create an array to hold all the text and structure information
            doc_structure = []
            for i, paragraph in enumerate(doc.paragraphs, start=1):
                # Check and update the section index and store section header if necessary
                if paragraph.style.name.startswith("Heading"):
                    current_section_index += 1
                    section_headers[current_section_index] = paragraph.text

                paragraph_info = {
                    "content": paragraph.text,
                    "location": {
                        "section_index": current_section_index,
                        "paragraph_index": i,
                        "section_header": section_headers.get(
                            current_section_index, ""
                        ),
                    },
                    "canSplit": True,
                }

                doc_structure.append(paragraph_info)

            return doc_structure

    def handle(self, file_content, key, split_params):
        with io.BytesIO(file_content) as f:
            doc = Document(f)

            # Create a structure to hold section header information
            current_section_index = 0
            section_headers = {}

            # Create an array to hold all the text and structure information
            doc_structure = []
            for i, paragraph in enumerate(doc.paragraphs, start=1):
                # Check and update the section index and store section header if necessary
                if paragraph.style.name.startswith("Heading"):
                    current_section_index += 1
                    section_headers[current_section_index] = paragraph.text

                paragraph_info = {
                    "text": paragraph.text,
                    "section_index": current_section_index,
                    "paragraph_index": i,
                    "section_header": section_headers.get(current_section_index, ""),
                }
                doc_structure.append(paragraph_info)

            full_text = "\n\n".join(p["text"] for p in doc_structure)
            chunks = self.intelligent_split(full_text, split_params)

            # Add additional DOCX-specific location data
            for chunk in chunks:
                chunk_location_info = self.find_location_in_doc_structure(
                    chunk["content"], doc_structure
                )
                chunk["location"].update(chunk_location_info)

            return chunks

    def find_location_in_doc_structure(self, chunk, doc_structure):
        location_info = dict()
        current_text = ""

        for item in doc_structure:
            if current_text and chunk.startswith(current_text):
                location_info = {
                    "section_index": item["section_index"],
                    "section_header": item["section_header"],
                    "paragraph_index": item["paragraph_index"],
                }
                break
            if item["text"]:
                current_text += item["text"] + "\n\n"

        return location_info
